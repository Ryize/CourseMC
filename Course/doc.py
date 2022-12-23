import datetime
import os
import uuid

import docx
from django.http import HttpResponse
from docx.shared import RGBColor
from htmldocx import HtmlToDocx


def save_report(text_report):
    directory = 'report'
    filename = 'Report_{unique_symbol}.txt'.format(unique_symbol=uuid.uuid4())
    if not os.path.exists(directory):
        os.mkdir(directory)
    response = HttpResponse(text_report, content_type='text/plain')
    response['Content-Disposition'] = 'attachment; filename={0}'.format(
        filename,
    )
    return response


def docx_worker(schedules):
    doc = docx.Document()
    doc.add_heading('Программа курса Python разработки', 0)
    paragraph = doc.add_paragraph('Получено: ')
    month_conver = {
        'Unknown': '-',
        'January': 'Января',
        'February': 'Февраля',
        'March': 'Марта',
        'April': 'Апреля',
        'May': 'Мая',
        'June': 'Июня',
        'July': 'Июля',
        'August': 'Августа',
        'September': 'Сентября',
        'October': 'Октября',
        'November': 'Ноября',
        'December': 'Декабря',
    }
    date = datetime.date.today().strftime('%d %B %Y').split()
    date[1] = month_conver[date[1]]
    run = paragraph.add_run('{date} года'.format(date=' '.join(date)))
    set_font_color(run)
    table_worker(doc, HtmlToDocx(), schedules)
    return doc


def table_worker(doc, new_parser, schedules):
    table = doc.add_table(rows=schedules.count(), cols=2)
    table.style = 'Table Grid'
    for row, schedule in enumerate(schedules):
        cell = table.cell(row, 0)
        new_parser.add_html_to_document(schedule.theme, cell)
        cell = table.cell(row, 1)
        new_parser.add_html_to_document(schedule.lesson_materials, cell)


def set_font_color(run):
    red = 255
    green = 10
    blue = 10
    run.font.color.rgb = RGBColor(red, green, blue)
