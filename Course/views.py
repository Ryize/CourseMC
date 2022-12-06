import datetime
import os
from io import BytesIO

import docx
from django.conf.urls import url
from django.contrib.auth import get_user
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.models import User
from django.http import HttpResponse, JsonResponse
from django.shortcuts import redirect, render
from django.views.generic import ListView
from django.views.generic.edit import FormView
from docx.shared import RGBColor
from htmldocx import HtmlToDocx

from reviews.models import Review

from .forms import *
from .models import LearnGroup, Schedule, Student, StudentQuestion


class StudentRecordView(FormView):
    template_name = "Course/index.html"
    form_class = StudentForm
    login_url = "/login/"

    def form_valid(self, form):
        form.save()
        name = form.cleaned_data["name"]
        email = form.cleaned_data["email"]
        password = form.cleaned_data["password"]
        user = User.objects.create_user(name, email, password)
        user.save()
        response = {
            "success": True,
        }
        return JsonResponse(response)

    def form_invalid(self, form):
        response = {"success": False, "error_message": "Форма заполнена не верно!"}
        return JsonResponse(response)

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context["reviews_count"] = Review.objects.all().count()
        return context


class TimetableView(LoginRequiredMixin, ListView):
    model = Schedule
    template_name = "Course/timetable.html"
    context_object_name = "schedules"
    paginate_by = 16
    queryset = Schedule.objects.filter(is_display=True)

    def get_queryset(self):
        student = Student.objects.filter(name=get_user(self.request).username).first()
        student_group = Schedule.objects.filter(group=student.groups, is_display=True).order_by(
            "-weekday"
        )
        theme = self.__get_param("theme")
        if theme:
            student_group = student_group.filter(theme__icontains=theme)
        if self.__get_param("lesson_type"):
            student_group = student_group.filter(
                lesson_type__icontains=self.__get_param("lesson_type")
            )
        if self.__get_param("absent"):
            student_group = student_group.filter(absent__name__iexact=f"{student.name}")
        return student_group

    def get_context_data(self, *, object_list=None, **kwargs):
        context = super().get_context_data(**kwargs)
        context["reviews_count"] = Review.objects.all().count()
        student = Student.objects.filter(name=get_user(self.request).username).first()
        context["create_report"] = Schedule.objects.filter(
            group=student.groups
        ).order_by("-weekday")
        return context

    def dispatch(self, request, *args, **kwargs):
        student = Student.objects.filter(name=get_user(self.request).username).first()
        if not student:
            return redirect("home")
        return super().dispatch(request, *args, **kwargs)

    def __get_param(self, name):
        return self.request.GET.get(name)


@login_required
def download_report(request, group_id):
    group = LearnGroup.objects.filter(pk=group_id).first()
    if not group:
        return redirect("home")
    schedules = group.schedules.all()
    students = group.students
    schedule_count, students_count = schedules.count(), students.count()
    schedules_primary = schedules.filter(lesson_type="Ключевой урок").count()
    schedules_new_theme = schedules.filter(lesson_type="Новая тема").count()
    practice_schedules = schedules.filter(lesson_type="Практика").count()

    result_data = f"""Отчёт о группе {group.title}
{'-' * 64}
Количество расписаний: {schedule_count}
Ключевых уроков: {schedules_primary}
Новых тем: {schedules_new_theme}
Практики: {practice_schedules}
Участников ({students_count}):
    """
    for student in students.all():
        result_data += f"""
    Имя: {student.name}
    Контакты: {student.contact or 'Не указаны!'}
    Пропущенных урока: {student.absents.count()}
        """
    result_data += f"\n\nЧасов обучения: {schedules.count() * 1.5}"

    directory = "report"
    filename = f"Report_{random.randint(10000, 999999999)}.txt"
    if not os.path.exists(directory):
        os.mkdir(directory)

    content = result_data
    response = HttpResponse(content, content_type="text/plain")
    response["Content-Disposition"] = "attachment; filename={0}".format(filename)
    return response


def get_training_program(response):
    schedules = (
        Schedule.objects.filter(
            group=LearnGroup.objects.filter(title="Вояджер").first()
        )
        .order_by("weekday")
        .all()
    )
    destination_document_file = BytesIO()
    doc = docx.Document()
    doc.add_heading(f"Программа курса Python разработки", 0)
    paragraph = doc.add_paragraph("Получено: ")
    month_conver = {
        "Unknown": "-",
        "January": "Января",
        "February": "Февраля",
        "March": "Марта",
        "April": "Апреля",
        "May": "Мая",
        "June": "Июня",
        "July": "Июля",
        "August": "Августа",
        "September": "Сентября",
        "October": "Октября",
        "November": "Ноября",
        "December": "Декабря",
    }
    date = datetime.date.today().strftime("%d %B %Y").split()
    date[1] = month_conver[date[1]]
    run = paragraph.add_run(f'{" ".join(date)} года')
    run.font.color.rgb = RGBColor(255, 10, 20)
    new_parser = HtmlToDocx()

    table = doc.add_table(rows=schedules.count(), cols=2)
    table.style = "Table Grid"

    for row, schedule in enumerate(schedules):
        cell = table.cell(row, 0)
        new_parser.add_html_to_document(schedule.theme, cell)
        cell = table.cell(row, 1)
        new_parser.add_html_to_document(schedule.lesson_materials, cell)

    if not os.path.exists("programCoursePython"):
        os.mkdir("programCoursePython")

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = "attachment; filename=programCoursePython.docx"
    doc.save(response)
    return response


@login_required
def ask_question(request):
    question = request.GET.get("question")
    response = {
        "is_taken": False,
    }
    if not question:
        return JsonResponse(response, status=412)
    student = Student.objects.filter(name=request.user.username).first()
    if not student:
        return JsonResponse(response, status=403)
    group = student.groups
    student_question = StudentQuestion(group=group, question=question)
    student_question.save()
    response["is_taken"] = True
    return JsonResponse(response, status=200)


@login_required
def get_filter_data(request):
    random_data = random.randint(0, 99999999999999999)
    result = f"{random_data}_Ты нашёл пасхалку, красавчик!!!"
    result_dict = {
        "answer": result,
        "code": "OK",
    }
    return JsonResponse(result_dict, status=200)


@login_required
def create_group(request):
    if not request.user.is_staff:
        return redirect('/')
    if request.method == 'GET':
        return render(request, 'Course/groups.html')
    title = request.POST.get("title")
    new_group = LearnGroup(title=title)
    new_group.save()
    all_plan_lessons = Schedule.objects.filter(group=LearnGroup.objects.get(id=3)).order_by('weekday')
    for i in all_plan_lessons:
        schedule = Schedule(group=new_group, theme=i.theme, weekday=i.weekday, time_lesson=i.time_lesson,
                            lesson_materials=i.lesson_materials, is_display=False)
        schedule.save()
    return redirect('/')
